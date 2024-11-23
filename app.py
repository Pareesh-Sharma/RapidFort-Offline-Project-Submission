# from flask import Flask, request, jsonify, render_template, send_file
# import os
# from werkzeug.utils import secure_filename
# from docx import Document
# import pdfkit
# from PyPDF2 import PdfReader, PdfWriter
# import tempfile

# app = Flask(__name__)
# app.config['UPLOAD_FOLDER'] = 'uploads'
# os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# # Home route
# @app.route('/')
# def index():
#     return render_template('index.html')

# # Upload file route
# @app.route('/upload', methods=['POST'])
# def upload_file():
#     if 'file' not in request.files:
#         return jsonify({'error': 'No file part'}), 400
#     file = request.files['file']
#     if file.filename == '':
#         return jsonify({'error': 'No selected file'}), 400

#     if file and file.filename.endswith('.docx'):
#         filename = secure_filename(file.filename)
#         file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
#         file.save(file_path)

#         # Extract metadata
#         doc = Document(file_path)
#         metadata = {
#             'paragraph_count': len(doc.paragraphs),
#             'filename': filename
#         }

#         # Convert to PDF
#         pdf_file_path = file_path.replace('.docx', '.pdf')
#         pdfkit.from_file(file_path, pdf_file_path)

#         # Respond with metadata and PDF download link
#         return jsonify({
#             'metadata': metadata,
#             'pdf_url': f'/download/{filename.replace(".docx", ".pdf")}'
#         })

#     return jsonify({'error': 'Invalid file format, only .docx allowed'}), 400

# # Download PDF route
# @app.route('/download/<filename>')
# def download_file(filename):
#     file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
#     if os.path.exists(file_path):
#         return send_file(file_path, as_attachment=True)
#     return jsonify({'error': 'File not found'}), 404

# if __name__ == '__main__':
#     app.run(debug=True)

from flask import Flask, request, jsonify, render_template, send_file
import os
from werkzeug.utils import secure_filename
from docx import Document
import pdfkit
import tempfile

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Home route
@app.route('/')
def index():
    return render_template('index.html')

# Helper function to convert .docx to HTML
def docx_to_html(docx_path):
    document = Document(docx_path)
    html_content = "<html><body>"
    for paragraph in document.paragraphs:
        html_content += f"<p>{paragraph.text}</p>"
    html_content += "</body></html>"
    return html_content

# Upload file route
@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400

    if file and file.filename.endswith('.docx'):
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)

        # Extract metadata
        try:
            doc = Document(file_path)
            metadata = {
                'paragraph_count': len(doc.paragraphs),
                'filename': filename
            }
        except Exception as e:
            return jsonify({'error': f'Failed to extract metadata: {str(e)}'}), 500

        # Convert to HTML
        try:
            html_content = docx_to_html(file_path)
            html_path = os.path.join(app.config['UPLOAD_FOLDER'], filename.replace('.docx', '.html'))
            with open(html_path, 'w') as html_file:
                html_file.write(html_content)
        except Exception as e:
            return jsonify({'error': f'Failed to convert .docx to HTML: {str(e)}'}), 500

        # Convert HTML to PDF
        try:
            pdf_file_path = file_path.replace('.docx', '.pdf')
            pdfkit.from_file(html_path, pdf_file_path)
        except Exception as e:
            return jsonify({'error': f'Failed to convert HTML to PDF: {str(e)}'}), 500

        # Respond with metadata and PDF download link
        return jsonify({
            'metadata': metadata,
            'pdf_url': f'/download/{os.path.basename(pdf_file_path)}'
        })

    return jsonify({'error': 'Invalid file format, only .docx allowed'}), 400

# Download PDF route
@app.route('/download/<filename>')
def download_file(filename):
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    if os.path.exists(file_path):
        return send_file(file_path, as_attachment=True)
    return jsonify({'error': 'File not found'}), 404

if __name__ == '__main__':
    app.run(debug=True)
